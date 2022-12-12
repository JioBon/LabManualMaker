# Imports
from multiprocessing import context
from tkinter import Widget
from xml.dom.minidom import Document
from django.shortcuts import render, redirect
from django.contrib.auth import authenticate, login, logout
from django.contrib import messages
from django.contrib.auth.decorators import login_required, user_passes_test
from django.core.exceptions import ObjectDoesNotExist
from django.db.models import Q
import requests

from .models import *
from .filters import LabFilter
from .forms import *

from docx import *
from PyRTF import document, Renderer
import os
from django.conf import settings
from django.http import HttpResponse, Http404
from mailmerge import MailMerge
from bs4 import BeautifulSoup
from html.parser import HTMLParser
from docxtpl import DocxTemplate, RichText
from htmldocx import HtmlToDocx


# Class for html parser
# class DocumentHTMLParser(HTMLParser):
#     def __init__(self, document):
#         HTMLParser.__init__(self)
#         self.document = document
#         self.paragraph  = None
#         self.run = None
    
#     def add_paragraph_and_feed(self, html):
#         self.paragraph = self.document.add_paragraph()
#         self.feed(html)

#     def handle_starttag(self, tag, attrs):
#         self.run = self.paragraph.add_run()

#         if tag in ["ul"]:
#             self.run.add_break()
#         if tag in ["li"]:
#             self.run.add_text(u'        \u2022    ')

#     def handle_endtag(self, tag):
#         if tag in ["li"]:
#             self.run.add_break()
    
#     def handle_data(self, data):
#         self.run.add_text(data)

# Create your views here.
# ----------------- Document Download ----------------- #

# def testTeacherDocument(request, pk):
#     user = request.user
#     labmanual = Lab_Manual.objects.get(id=pk)

#     template = os.path.join(settings.TEMPLATE_ROOT, 'labtemplate.docx')
#     docxdocument = DocxTemplate(template)
#     desc_document = Document()
#     new_parser = HtmlToDocx()
    
#     new_parser.add_html_to_document(labmanual.objectives, desc_document)
#     desc_document.save(os.path.join(settings.TEMPLATE_ROOT, f'subdoc.docx'))

#     # doc_html_parser = DocumentHTMLParser(docxdocument)
#     # test = doc_html_parser.add_paragraph_and_feed(labmanual.objectives)
#     # print(test)

#     sub_doc = docxdocument.new_subdoc(os.path.join(settings.TEMPLATE_ROOT, f'subdoc.docx'))

#     context = {
#         'objectives': sub_doc,

#     }


#     docxdocument.render(context)
#     docxdocument.save(os.path.join(settings.TEMPLATE_ROOT, f'{labmanual.lab_name}.docx'))

    
#     # docxdocument.render(context)

#     # docxdocument.write(os.path.join(settings.TEMPLATE_ROOT, f'{labmanual.lab_name}.docx'))

#     file_dir = os.path.join(settings.TEMPLATE_ROOT, f'{labmanual.lab_name}.docx')
#     if os.path.exists(file_dir):
#         with open(file_dir, 'rb') as fh:
#             response = HttpResponse(fh.read(), content_type="application/vnd.ms-excel")
#             response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file_dir)
#             print('gumana')
#             return response

#     # Renderer.Renderer.Write(os.path.join(settings.TEMPLATE_ROOT, f'{labmanual.lab_name}.rtf'))

#     return redirect('mainsite:viewlab', pk=pk)

def TeacherDocument(request, pk):
    user = request.user
    labmanual = Lab_Manual.objects.get(id=pk)

    template = os.path.join(settings.TEMPLATE_ROOT, 'labtemplate.docx')
    docxdocument = MailMerge(template)

    docxdocument.merge(
    title=f'{labmanual.activity_name}',
    actno=f'{labmanual.activity_no}',
    coursecode=f'{labmanual.course.course.course_code}',
    coursetitle=f'{labmanual.course.course.course_title}',
    instructor=f'{labmanual.instructor}',
    objectives=f'{labmanual.objectives}',
    ilo=f'{labmanual.ilo}',
    discussion=f'{labmanual.discussion}',
    resources=f'{labmanual.resources}',
    procedure=f'{labmanual.procedure}',
    results=f'{labmanual.results}',
    supplementary=f'{labmanual.supplementary}',
    observation=f'{labmanual.observation}',
    conclusion=f'{labmanual.conclusion}'
)
    docxdocument.write(os.path.join(settings.TEMPLATE_ROOT, f'{labmanual.lab_name}.docx'))

    file_dir = os.path.join(settings.TEMPLATE_ROOT, f'{labmanual.lab_name}.docx')
    if os.path.exists(file_dir):
        with open(file_dir, 'rb') as fh:
            response = HttpResponse(fh.read(), content_type="application/vnd.ms-excel")
            response['Content-Disposition'] = 'inline; filename=' + os.path.basename(file_dir)
            print('gumana')
            return response

    Renderer.Renderer.Write(os.path.join(settings.TEMPLATE_ROOT, f'{labmanual.lab_name}.rtf'))

    return redirect('mainsite:viewlab', pk=pk)

# ----------------- Authentication ----------------- #

def loginpage(request):
    if request.user.is_authenticated:
        return redirect('mainsite:home')
    else:
        reg_form = UserAdminCreationForm()
        if 'login-submit' in request.POST:
            if request.method == 'POST':
                email = request.POST.get("email")
                password = request.POST.get("password")

                user = authenticate(request, email=email, password=password)

                if user is not None:
                    login(request, user)
                    return redirect('mainsite:home')
                    
                else:
                    messages.info(request,'Username OR Password is incorrect')
        elif 'register-submit' in request.POST:
            if request.method == 'POST':
                regform = UserAdminCreationForm(request.POST)
                if regform.is_valid():
                    regform.save()
                    messages.success(request, 'Successfully created an account')
                else:
                    messages.success(request, 'Failed to created an account')
                
    context = {'reg_form': reg_form}
    return render(request, "main/authenticate.html", context)

@login_required(login_url='mainsite:login')
def logoutUser(request):
    logout(request)
    return redirect('mainsite:login')

# ----------------- Home ----------------- #

@login_required(login_url='mainsite:login')
def home(request):
    user = request.user
    if request.method == 'POST':
        searched = request.POST['searched']

        return redirect('mainsite:search', pk=searched)
    else:
        labmanuals = Lab_Manual.objects.filter(instructor=user)
        context = {'labmanuals': labmanuals}
        return render(request, "main/home.html", context)

@login_required(login_url='mainsite:login')
def about_us(request):
    return render(request, "main/aboutus.html")

@login_required(login_url='mainsite:login')
def contact(request):
    return render(request, "main/contact.html")

@login_required(login_url='mainsite:login')
def searchpage(request, pk):
    user = request.user
    # pk = pk.split()
    sharinglist = Sharing.objects.filter(instructor=user)
    shared=[]
    # sharinglist = Sharing.objects.all()
    
    if request.method == 'POST':
        searched = request.POST['searched']

        return redirect('mainsite:search', pk=searched)


    labmanuals = Lab_Manual.objects.filter( (Q(lab_name__icontains=pk) | Q(activity_name__icontains=pk) | 
        Q(course__course__course_code__icontains=pk) | 
        Q(course__course__course_title__icontains=pk) | 
        Q(instructor__first_name__icontains=pk) | 
        Q(instructor__last_name__icontains=pk)))

    for field in sharinglist:
        shared.append(field.LabManual)

    return render(request, "main/search.html", {
        'searched': pk,
        'labmanuals': labmanuals,
        'sharinglist': shared,
    })

# ----------------- Profile ----------------- #

@login_required(login_url='mainsite:login')
def profile(request):
    user = request.user
    courselist = CourseInstructor.objects.filter(instructor=user)

    for i in courselist:
        print(i.course.course_code, request.POST)
        if i.course.course_code in request.POST:
            delcourse = CourseInstructor.objects.get(instructor=user, course=i.course)
            delcourse.delete()
            messages.success(request, 'Course successfully Deleted')
    
    if 'eprofile' in request.POST:
        return redirect('mainsite:editprofile')

    context = {"user": user, "courses": CourseInstructor.objects.filter(instructor=user)}
    return render(request, "main/profile.html", context)

@login_required(login_url='mainsite:login')
def editProfile(request):
    user = request.user
    
    form = EditProfile(instance=user)

    context = {"user": user, "form": form}
    return render(request, "main/editprofile.html", context)

@login_required(login_url='mainsite:login')
def editCourse(request):
    user = request.user

    courselist = CourseInstructor.objects.filter(instructor=user)
    courseform = NewCourse(initial={'instructor': user})

    if 'addcourse' in request.POST:
        form = NewCourse(request.POST)
        getcourse = request.POST.get("course")
        if form.is_valid():
            try:
                checkCourse = CourseInstructor.objects.get(instructor=user, course=getcourse)
            except ObjectDoesNotExist:
                form.save()

            messages.success(request, 'Course Already Exist to this user')

    for i in courselist:
        if i.course.course_code in request.POST:
            delcourse = CourseInstructor.objects.get(instructor=user, course=i.course)
            delcourse.delete()
            messages.success(request, 'Course successfully Deleted')

    context = {"user": user, "courses": CourseInstructor.objects.filter(instructor=user) ,"form": courseform}
    return render(request, "main/editcourse.html", context)

# ----------------- Laboratory ----------------- #

@login_required(login_url='mainsite:login')
def createlab(request):
    user = request.user
    usercourselist = CourseInstructor.objects.filter(instructor=user)

    form = NewLab(initial={'instructor': user})
    form.fields['course'].queryset = usercourselist
    if 'newlab' in request.POST:
        form = NewLab(request.POST,
            request.FILES, initial={'instructor': user})
        if form.is_valid():
            form.save()
            return redirect('mainsite:home')

    return render(request, 'main/createlab.html', {
        'form': form,
        })

@login_required(login_url='mainsite:login')
def viewlab(request, pk):
    user = request.user
    labmanual = Lab_Manual.objects.get(id=pk)
    uri = request.build_absolute_uri()
    # print(uri)
    # page = requests.get(uri)
    # soup = BeautifulSoup(uri, "html.parser")
    # print(soup)
    # print("uri: ", uri)
    # print("soup", soup)
    # print("rtf data:", request.POST)
    # test = soup.find("p")
    if 'editlab' in request.POST:
        return redirect('mainsite:editlab', pk=pk)
    if 'teacherdownload' in request.POST:
        return redirect('mainsite:teacherdownload', pk=pk)


    context = {'labelements': labmanual}
    return render(request, 'main/viewlab.html', context)

@login_required(login_url='mainsite:login')
def sharedlab(request, pk):
    user = request.user
    labmanual = Lab_Manual.objects.get(id=pk)
    shareform = ShareLab(initial={'LabManual':labmanual})
    sharelist = Sharing.objects.filter(LabManual=labmanual)

    if 'New-Share-Submit' in request.POST:
        form = ShareLab(request.POST)
        print(request.POST.get("instructor"))
        if form.is_valid():
            try:
                checkCourse = Sharing.objects.get(instructor=request.POST.get("instructor"))
            except ObjectDoesNotExist:
                form.save()

    for i in sharelist:
        if i.instructor.__str__() in request.POST:
            delcourse = Sharing.objects.get(instructor=i.instructor)
            delcourse.delete()
            messages.success(request, 'Course successfully Deleted')
            sharelist = Sharing.objects.filter(LabManual=labmanual)
    
    context = {
        'labelements': labmanual, 'sharingform': shareform, 'sharelist': sharelist
    }
    if labmanual.instructor == user:
        return render(request, 'main/sharing.html', context)
    else:
        pass


@login_required(login_url='mainsite:login')
def editlab(request, pk):
    user = request.user
    labmanual = Lab_Manual.objects.get(id=pk)
    if labmanual.instructor == request.user:
        form = NewLab(instance=labmanual)
        if request.method == 'POST':
            form = form = NewLab(request.POST,
            request.FILES, instance=labmanual)
            if form.is_valid():
                form.save()
                return redirect('mainsite:viewlab', pk=pk)
        context = {'form': form, 'labelements': labmanual}
        return render(request, 'main/createlab.html', context )
    else:
        redirect('mainsite:home')

    